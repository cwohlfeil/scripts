import docx
doc = docx.Document('.\\demo.docx')
len(doc.paragraphs)
doc.paragraphs[0].text
doc.paragraphs[1].text
len(doc.paragraphs[1].runs)
doc.paragraphs[1].runs[0].text
doc.paragraphs[1].runs[1].text

doc = docx.Document()
doc.add_paragraph('Hello world!')
doc.add_heading('Header 0', 0)
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.add_picture('.\\zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
doc.save('.\\helloworld.docx')


def get_text(filename):
    doc2 = docx.Document(filename)
    full_text = []
    for para in doc2.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

print(get_text('.\\demo.docx'))


